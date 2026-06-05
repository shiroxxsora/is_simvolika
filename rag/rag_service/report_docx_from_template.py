from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from docx.shared import Cm

from rag_service.report_models import ReportPhoto

logger = logging.getLogger(__name__)


class TemplateFillError(RuntimeError):
    pass


def _replace_paragraph_text_exact(doc: Document, old: str, new: str) -> bool:
    """Replace paragraph text only when it matches exactly; preserves paragraph style."""
    for p in doc.paragraphs:
        if (p.text or "").strip() == old.strip():
            # Clear runs, then set a single run; paragraph style remains.
            for r in list(p.runs):
                try:
                    r._element.getparent().remove(r._element)  # noqa: SLF001
                except Exception:
                    pass
            p.add_run(new)
            return True
    return False


def _replace_paragraph_text_startswith(doc: Document, prefix: str, new: str) -> bool:
    """Replace paragraph text when it starts with prefix; preserves paragraph style."""
    pfx = (prefix or "").strip()
    if not pfx:
        return False
    for p in doc.paragraphs:
        cur = (p.text or "").strip()
        if cur.startswith(pfx):
            for r in list(p.runs):
                try:
                    r._element.getparent().remove(r._element)  # noqa: SLF001
                except Exception:
                    pass
            p.add_run(new)
            return True
    return False


def _clear_block_after_anchor_until(
    doc: Document, anchor_prefix: str, stop_before_prefix: str, *, max_remove: int = 200
) -> bool:
    """Remove all paragraphs right after the anchor line until a paragraph with stop_before_prefix.

    The base template may contain a full example bibliography (numbered and long-form); all of it
    must go before we insert the generated numbered list, otherwise the document duplicates sources.
    """
    pfx = (anchor_prefix or "").strip()
    stop = (stop_before_prefix or "").strip()
    if not pfx or not stop:
        return False
    for idx, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t.startswith(pfx):
            continue
        n = 0
        found_stop = False
        while n < max_remove:
            body = list(doc.paragraphs)
            if idx + 1 >= len(body):
                break
            nxt = body[idx + 1]
            t2 = (nxt.text or "").strip()
            if t2.startswith(stop):
                found_stop = True
                break
            try:
                nxt._element.getparent().remove(nxt._element)  # noqa: SLF001
            except Exception:  # noqa: BLE001
                logger.warning("Template DOCX: failed to remove paragraph after source anchor")
                return False
            n += 1
        if not found_stop:
            logger.warning(
                "Template DOCX: end marker not found for source block (expected %r before limit)",
                stop,
            )
            return False
        return True
    return False


def _paragraph_style_name_for_insert(doc: Document, anchor_prefix: str) -> str | None:
    """Стилистика первой строки старого списка (до удаления) — для вставляемого нумерованного списка."""
    pfx = anchor_prefix.strip()
    ps = list(doc.paragraphs)
    for i, p in enumerate(ps):
        t = (p.text or "").strip()
        if not t.startswith(pfx):
            continue
        if i + 1 < len(ps) and (ps[i + 1].text or "").strip():
            s = ps[i + 1].style
        else:
            s = p.style
        if s and s.name:
            return s.name
    return None


def _remove_all_inline_pictures_from_document(doc: Document) -> int:
    """Удаляет встроенные в шаблон рисунки (пример в «Исследование»); новое фото вставляется по маркеру."""
    removed = 0

    def _strip_run(run) -> None:
        nonlocal removed
        r = run._element  # noqa: SLF001
        to_remove = [el for el in r.iter() if (el.tag or "").endswith(("}drawing", "}pict"))]
        for el in to_remove:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                removed += 1

    for p in doc.paragraphs:
        for run in p.runs:
            _strip_run(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _strip_run(run)
    return removed


def _rebuild_research_and_photo_marker(
    doc: Document,
    research_paragraphs: list[str],
    *,
    section_title: str = "Исследование",
    end_before_starts: str = "Вывод",
    photo_marker: str = "(Фото № 1)",
) -> bool:
    """Между заголовком «Исследование» и «Вывод» вставить текст, затем маркер для фото (шаблонные абзацы и рисунок удаляются)."""
    ps = list(doc.paragraphs)
    t_idx: int | None = None
    e_idx: int | None = None
    for i, p in enumerate(ps):
        if (p.text or "").strip() == section_title:
            t_idx = i
            continue
        if t_idx is not None and (p.text or "").strip().startswith(end_before_starts):
            e_idx = i
            break
    if t_idx is None or e_idx is None:
        logger.warning("Template DOCX: research block anchors not found (title/Вывод)")
        return False
    ref_style = ps[t_idx + 1].style if t_idx + 1 < e_idx and (ps[t_idx + 1].text or "").strip() else None
    if ref_style is None or not (ref_style.name or "").strip():
        ref_style = ps[t_idx].style
    name = (ref_style.name or "").strip() if ref_style and ref_style.name else None
    n = 0
    while n < 200 and t_idx + 1 < len(list(doc.paragraphs)):
        p = list(doc.paragraphs)[t_idx + 1]
        if (p.text or "").strip().startswith(end_before_starts):
            break
        try:
            p._element.getparent().remove(p._element)  # noqa: SLF001
        except Exception:  # noqa: BLE001
            logger.warning("Template DOCX: failed to clear research template block")
            return False
        n += 1
    anchor_p = list(doc.paragraphs)[t_idx]
    last = anchor_p._p  # noqa: SLF001
    for line in research_paragraphs:
        if not (line or "").strip():
            continue
        new_p = doc.add_paragraph(line.strip())
        if name:
            try:
                new_p.style = name
            except Exception:  # noqa: BLE001
                pass
        last.addnext(new_p._p)  # noqa: SLF001
        last = new_p._p  # noqa: SLF001
    m_p = doc.add_paragraph(photo_marker)
    if name:
        try:
            m_p.style = name
        except Exception:  # noqa: BLE001
            pass
    last.addnext(m_p._p)  # noqa: SLF001
    return True


def _insert_sources_after_anchor(
    doc: Document,
    anchor_prefix: str,
    sources_lines: list[str],
    *,
    stop_before_prefix: str = "Содержание исследования",
) -> bool:
    """Find paragraph starting with anchor_prefix, drop template example up to the next section, then insert lines."""
    anchor_pfx = anchor_prefix.strip()
    ref_style = _paragraph_style_name_for_insert(doc, anchor_pfx)
    if not _clear_block_after_anchor_until(doc, anchor_pfx, stop_before_prefix):
        return False
    if not sources_lines:
        return True
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith(anchor_pfx):
            last = p._p  # noqa: SLF001
            for line in sources_lines:
                new_p = doc.add_paragraph(line)
                if ref_style:
                    try:
                        new_p.style = ref_style
                    except Exception:  # noqa: BLE001
                        pass
                last.addnext(new_p._p)  # noqa: SLF001
                last = new_p._p  # noqa: SLF001
            return True
    return False


def _insert_photo_after_marker(doc: Document, marker_text: str, photo: ReportPhoto) -> bool:
    """Insert an image after a paragraph whose text matches marker_text (e.g. '(Фото № 1)')."""
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == marker_text.strip():
            # Insert a new paragraph after marker.
            # python-docx does not expose insert-after directly, so we add at end and move XML.
            new_p = doc.add_paragraph()
            new_p.alignment = p.alignment
            run = new_p.add_run()
            try:
                run.add_picture(io.BytesIO(photo.image_bytes), width=Cm(12.5))
            except Exception as exc:
                logger.warning("Template DOCX: failed to embed photo: %s", exc)
                return False

            p_element = p._p  # noqa: SLF001
            new_element = new_p._p  # noqa: SLF001
            p_element.addnext(new_element)
            return True
    return False


def build_docx_from_base_template(
    *,
    base_template_path: Path,
    replacements_exact: dict[str, str],
    replacements_prefix: dict[str, str] | None = None,
    sources_anchor_prefix: str | None = None,
    sources_lines: list[str] | None = None,
    research_paragraphs: list[str] | None = None,
    photos: list[ReportPhoto],
) -> bytes:
    """Загрузка шаблона: убрать встроенные рисунки, подстановка текста, список источников, блок исследования, фото."""
    doc = Document(str(base_template_path))
    n_img = _remove_all_inline_pictures_from_document(doc)
    if n_img:
        logger.info("Template DOCX: removed %s embedded image(s) from example layout", n_img)

    missing: list[str] = []
    for old, new in replacements_exact.items():
        ok = _replace_paragraph_text_exact(doc, old, new)
        if not ok:
            missing.append(old)

    prefix_missing: list[str] = []
    for old, new in (replacements_prefix or {}).items():
        ok = _replace_paragraph_text_startswith(doc, old, new)
        if not ok:
            prefix_missing.append(old)

    if sources_anchor_prefix and sources_lines is not None:
        ok = _insert_sources_after_anchor(doc, sources_anchor_prefix, sources_lines)
        if not ok:
            logger.warning("Template DOCX: sources step failed, anchor=%r", sources_anchor_prefix)

    if research_paragraphs is not None:
        ok = _rebuild_research_and_photo_marker(doc, research_paragraphs)
        if not ok:
            logger.warning("Template DOCX: research block rebuild failed")

    for photo in photos:
        marker = f"(Фото № {photo.number})"
        ok = _insert_photo_after_marker(doc, marker, photo)
        if not ok:
            logger.warning("Template DOCX: photo marker not found: %r", marker)

    out = io.BytesIO()
    doc.save(out)
    if missing:
        logger.warning("Template DOCX: %s replacement anchors not found", len(missing))
    if prefix_missing:
        logger.warning("Template DOCX: %s prefix anchors not found", len(prefix_missing))
    return out.getvalue()

