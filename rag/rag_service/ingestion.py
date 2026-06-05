import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

import fitz
from docx import Document as DocxDocument
from PIL import Image, ImageFilter, ImageOps
from pypdf import PdfReader
import pytesseract

from rag_service.source_page_norm import normalize_source_page


@dataclass(frozen=True)
class SourceMeta:
    source_doc: str | None
    source_chapter: str | None
    source_page: str | None


@dataclass(frozen=True)
class IngestDocument:
    doc_id: str
    doc_name: str
    content: str
    meta: SourceMeta
    units: list["TextUnit"]
    illustration_segments: tuple[tuple[str, str], ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class TextUnit:
    text: str
    chapter: str | None
    page: str | None


@dataclass(frozen=True)
class ChunkRecord:
    text: str
    chapter: str | None
    page: str | None


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

# Кеш prepare (scripts/prepare_ocr_docs.py): служебные файлы не должны попадать в RAG как документы.
_PREPARE_PAGE_CACHE_DIR = ".prepare_page_cache"


def _is_prepare_page_cache_path(path: Path) -> bool:
    return _PREPARE_PAGE_CACHE_DIR in path.parts
DOCX_VIRTUAL_PAGE_CHARS = 2200
PDF_OCR_FALLBACK_ENABLED = os.getenv("PDF_OCR_FALLBACK_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
def _tesseract_lang() -> str:
    return os.getenv("TESSERACT_LANG") or os.getenv("PDF_OCR_LANG") or "rus+eng"
PDF_OCR_FALLBACK_MODE = os.getenv("PDF_OCR_FALLBACK_MODE", "empty_only").strip().lower()


def _extract_meta(content: str, key: str) -> str | None:
    found = re.search(rf"(?im)^{re.escape(key)}:\s*(.+?)\s*$", content)
    if not found:
        return None
    value = found.group(1).strip()
    return value or None


def _strip_meta_lines(content: str) -> str:
    cleaned_lines: list[str] = []
    for line in content.splitlines():
        if re.match(r"(?i)^\s*source_(doc|chapter|page)\s*:", line):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines).strip()
    return re.sub(r"\n{3,}", "\n\n", cleaned)


def _extract_page_marker(line: str) -> str | None:
    match = re.match(r"^\s*\[PAGE\s+(\d+)\]\s*$", line, flags=re.IGNORECASE)
    if match:
        return normalize_source_page(match.group(1))
    return None


def _normalize_inline_page_markers(content: str) -> str:
    """VL иногда вставляет [PAGE N] в конец длинной строки; парсер видит только маркер на отдельной строке."""
    if "[PAGE" not in content.upper():
        return content
    text = content
    # Подряд [PAGE 1][PAGE 2] без перевода строки
    prev = None
    while prev != text:
        prev = text
        text = re.sub(
            r"(\[PAGE\s+\d+\])\s*(\[PAGE\s+\d+\])",
            r"\1\n\2",
            text,
            flags=re.IGNORECASE,
        )
    # Текст ... [PAGE 26]
    text = re.sub(r"([^\n])\s*(\[PAGE\s+\d+\])", r"\1\n\2", text, flags=re.IGNORECASE)
    # [PAGE 26] продолжение на той же строке
    text = re.sub(r"(\[PAGE\s+\d+\])\s+([^\n])", r"\1\n\2", text, flags=re.IGNORECASE)
    return text


def _is_illustration_rubric_line(line: str) -> bool:
    """Строки вроде «231 [ИЛЛЮСТРАЦИИ]» / «[ИЛЛЮСТРАЦИЯ 1]» — не заголовки глав, а метки VL."""
    text = line.strip()
    if len(text) > 200:
        return False
    if not re.search(r"ИЛЛЮСТРАЦ", text, flags=re.IGNORECASE):
        return False
    # Длинное предложение про иллюстрации — оставляем как обычный текст
    if len(text) > 90 and (". " in text or text.count(",") >= 2):
        return False
    return bool(
        re.fullmatch(
            r"\s*(?:\d{1,4}\s+)?\[?\s*(ИЛЛЮСТРАЦИИ|ИЛЛЮСТРАЦИЯ)\s*(\d+)?\s*\]?\s*",
            text,
            flags=re.IGNORECASE,
        )
    )


def _is_prison_case_rubric_line(line: str) -> bool:
    """Короткие строки учёта вроде «21 20. ИТК-5. 60-е гг.» — не заголовки разделов."""
    t = line.strip()
    if len(t) > 130:
        return False
    if re.match(r"^\d{1,3}\s+\d{1,3}\.\s", t):
        return True
    if len(t) < 100 and re.search(r"\bИТК\b", t, flags=re.IGNORECASE) and re.match(r"^\d", t):
        return True
    return False


def _remove_vl_illustration_description_blocks(text: str) -> str:
    """Убирает блоки VL «[ИЛЛЮСТРАЦИИ]» / «[ИЛЛЮСТРАЦИЯ N] …» до следующей такой метки; смысл для RAG в multimodal-страницах."""
    if "ИЛЛЮСТРАЦ" not in text.upper():
        return text
    cleaned = re.sub(
        r"\[(?:ИЛЛЮСТРАЦИИ|ИЛЛЮСТРАЦИЯ\s*\d+)\][^\[]*",
        " ",
        text,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r" {2,}", " ", cleaned)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def split_vl_page_to_body_and_illustrations(page_text: str) -> tuple[str, str]:
    """Делит текст страницы VL на тело и блок описаний иллюстраций (для кеша и split-выхода prepare)."""
    raw = page_text.strip()
    if not raw:
        return "", ""
    if "ИЛЛЮСТРАЦ" not in raw.upper():
        return raw, ""
    body = _remove_vl_illustration_description_blocks(raw)
    m = re.search(
        r"(\[(?:ИЛЛЮСТРАЦИИ|ИЛЛЮСТРАЦИЯ\s*\d+)\].*)",
        raw,
        flags=re.IGNORECASE | re.DOTALL,
    )
    ill = m.group(1).strip() if m else ""
    return body.strip(), ill


def extract_illustration_segments_from_stripped_content(stripped: str) -> list[tuple[str, str]]:
    """Связывает блоки [ИЛЛЮСТРАЦИИ]… с последним встреченным [PAGE n] (legacy-однопотоковый TXT)."""
    if "ИЛЛЮСТРАЦ" not in stripped.upper():
        return []
    segments: list[tuple[str, str]] = []
    current_page: str | None = None
    i = 0
    lines = stripped.splitlines()
    while i < len(lines):
        line = lines[i]
        m_page = re.match(r"^\s*\[PAGE\s+(\d+)\]\s*$", line, flags=re.IGNORECASE)
        if m_page:
            current_page = normalize_source_page(m_page.group(1))
            i += 1
            continue
        if re.search(r"\[(?:ИЛЛЮСТРАЦИИ|ИЛЛЮСТРАЦИЯ\s*\d+)\]", line, flags=re.IGNORECASE):
            buf = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if re.match(r"^\s*\[PAGE\s+", nxt, flags=re.IGNORECASE):
                    break
                if re.match(r"^\s*\[(?:ИЛЛЮСТРАЦИИ|ИЛЛЮСТРАЦИЯ\s*\d+)\]", nxt, flags=re.IGNORECASE):
                    break
                buf.append(nxt)
                i += 1
            block = "\n".join(buf).strip()
            if current_page and block:
                segments.append((current_page, block))
            continue
        i += 1
    return segments


def parse_prepare_split_format(meta_stripped: str) -> tuple[str | None, list[tuple[str, str]]]:
    """Разбор TXT с секциями ### RAG_PAGE_BODY / ### RAG_PAGE_ILLUSTRATIONS (см. prepare_ocr_docs --no-legacy-single-stream)."""
    if "### RAG_PAGE_BODY" not in meta_stripped:
        return None, []
    lines = meta_stripped.splitlines()
    body_out: list[str] = []
    ill_segments: list[tuple[str, str]] = []
    i = 0
    current_page: str | None = None

    while i < len(lines):
        line = lines[i]
        m_page = re.match(r"^\s*\[PAGE\s+(\d+)\]\s*$", line, flags=re.IGNORECASE)
        if m_page:
            page_raw = m_page.group(1)
            current_page = normalize_source_page(page_raw)
            body_out.append(f"[PAGE {page_raw}]")
            i += 1
            while i < len(lines) and lines[i].strip() != "### RAG_PAGE_BODY":
                if re.match(r"^\s*\[PAGE\s+", lines[i], flags=re.IGNORECASE):
                    return None, []
                body_out.append(lines[i])
                i += 1
            if i >= len(lines) or lines[i].strip() != "### RAG_PAGE_BODY":
                return None, []
            i += 1
            while i < len(lines):
                if lines[i].strip() == "### RAG_PAGE_ILLUSTRATIONS":
                    break
                if re.match(r"^\s*\[PAGE\s+", lines[i], flags=re.IGNORECASE):
                    break
                body_out.append(lines[i])
                i += 1
            while body_out and not body_out[-1].strip():
                body_out.pop()
            body_out.append("")
            if i < len(lines) and lines[i].strip() == "### RAG_PAGE_ILLUSTRATIONS":
                i += 1
                ill_lines: list[str] = []
                while i < len(lines) and not re.match(r"^\s*\[PAGE\s+", lines[i], flags=re.IGNORECASE):
                    ill_lines.append(lines[i])
                    i += 1
                ill_text = "\n".join(ill_lines).strip()
                if current_page and ill_text:
                    ill_segments.append((current_page, ill_text))
            continue
        if line.strip() in ("### RAG_PAGE_BODY", "### RAG_PAGE_ILLUSTRATIONS"):
            i += 1
            continue
        i += 1

    body_text = "\n".join(body_out).strip()
    body_text = re.sub(r"\n{3,}", "\n\n", body_text)
    return body_text, ill_segments


def _is_heading(line: str) -> bool:
    text = line.strip()
    if not text:
        return False
    if _is_illustration_rubric_line(line):
        return False
    if _is_prison_case_rubric_line(line):
        return False
    if len(text) > 120:
        return False
    if re.match(r"^(глава|раздел|часть|chapter|section)\s+[A-Za-zА-Яа-я0-9IVXLCM\-\.]+", text, flags=re.IGNORECASE):
        return True
    # «21 20. …» не считать заголовком (уже отфильтровано _is_prison_case_rubric_line)
    if re.match(r"^\d+(\.\d+)*\.?\s+\S+", text):
        if re.match(r"^\d{1,3}\s+\d{1,3}\.\s", text):
            return False
        return True
    alpha_chars = re.sub(r"[^A-Za-zА-Яа-я]", "", text)
    # Короткие капс-вставки («СВЯЛОЙ») — не заголовки; настоящие — обычно длиннее.
    if alpha_chars and text.isupper() and len(alpha_chars) >= 9:
        return True
    return False


def _extract_units(content: str, default_chapter: str | None, default_page: str | None) -> list[TextUnit]:
    current_chapter = default_chapter
    current_page = default_page
    units: list[TextUnit] = []
    paragraph_lines: list[str] = []

    def _flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = " ".join(line.strip() for line in paragraph_lines).strip()
        paragraph_lines.clear()
        if text:
            units.append(TextUnit(text=text, chapter=current_chapter, page=current_page))

    for raw_line in content.splitlines():
        line = raw_line.strip()
        page_marker = _extract_page_marker(line)
        if page_marker is not None:
            _flush_paragraph()
            current_page = page_marker
            continue

        if not line:
            _flush_paragraph()
            continue

        paragraph_lines.append(line)

    _flush_paragraph()
    return units


def _merge_rubric_only_units(units: list[TextUnit]) -> list[TextUnit]:
    """Склеивает отдельный короткий юнит «N [ИЛЛЮСТРАЦИИ]» со следующим на той же странице."""
    if len(units) < 2:
        return units
    merged: list[TextUnit] = []
    i = 0
    while i < len(units):
        cur = units[i]
        if (
            i + 1 < len(units)
            and _is_illustration_rubric_line(cur.text)
            and len(cur.text.strip()) < 160
            and (cur.page or "") == (units[i + 1].page or "")
        ):
            nxt = units[i + 1]
            combined = f"{cur.text.strip()}\n\n{nxt.text.strip()}"
            merged.append(
                TextUnit(text=combined, chapter=nxt.chapter, page=nxt.page)
            )
            i += 2
            continue
        merged.append(cur)
        i += 1
    return merged


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def iter_pdf_pages(path: Path, include_empty: bool = False) -> Iterator[tuple[int, str]]:
    reader = PdfReader(str(path))
    pdf_doc = None
    if PDF_OCR_FALLBACK_ENABLED:
        try:
            pdf_doc = fitz.open(str(path))
        except Exception:
            pdf_doc = None

    try:
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if PDF_OCR_FALLBACK_ENABLED and _needs_ocr_fallback(text):
                ocr_text = _ocr_pdf_page(pdf_doc, index - 1)
                if ocr_text:
                    text = ocr_text

            if include_empty or text:
                yield index, text
    finally:
        if pdf_doc is not None:
            pdf_doc.close()


def _read_pdf_pages(path: Path) -> list[tuple[int, str]]:
    return list(iter_pdf_pages(path, include_empty=False))


def _read_pdf(path: Path) -> str:
    pages = _read_pdf_pages(path)
    return "\n\n".join(f"[PAGE {index}]\n{text}" for index, text in pages)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines: list[str] = []
    virtual_page = 1
    current_page_chars = 0

    def _start_new_page() -> None:
        nonlocal virtual_page, current_page_chars
        virtual_page += 1
        current_page_chars = 0
        lines.append(f"[PAGE {virtual_page}]")

    lines.append(f"[PAGE {virtual_page}]")
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        is_heading_style = style_name.startswith("heading")
        has_page_break = _paragraph_has_page_break(paragraph)

        if has_page_break and current_page_chars > 0:
            _start_new_page()

        if text:
            lines.append(text)
            current_page_chars += len(text)
            if is_heading_style:
                # Blank line after headings helps unit separation.
                lines.append("")

        if current_page_chars >= DOCX_VIRTUAL_PAGE_CHARS:
            _start_new_page()

    return "\n".join(lines)


def _paragraph_has_page_break(paragraph) -> bool:
    for run in paragraph.runs:
        try:
            if run._element.xpath(".//w:br[@w:type='page']"):  # noqa: SLF001
                return True
        except Exception:
            continue
    return False


def _needs_ocr_fallback(text: str) -> bool:
    cleaned = text.strip()
    if PDF_OCR_FALLBACK_MODE == "off":
        return False
    if PDF_OCR_FALLBACK_MODE == "empty_only":
        return len(cleaned) == 0

    if len(cleaned) < 80:
        return True

    cyr = len(re.findall(r"[А-Яа-яЁё]", cleaned))
    lat = len(re.findall(r"[A-Za-z]", cleaned))

    # For Russian corpora, extracted text with no Cyrillic and many latin chars is likely broken.
    if cyr == 0 and lat > 120:
        return True

    # Too many mixed-script tokens is another signal of bad extraction.
    tokens = re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", cleaned)
    mixed_tokens = 0
    for token in tokens:
        has_lat = bool(re.search(r"[A-Za-z]", token))
        has_cyr = bool(re.search(r"[А-Яа-яЁё]", token))
        if has_lat and has_cyr:
            mixed_tokens += 1
    if tokens and (mixed_tokens / len(tokens)) > 0.2:
        return True

    return False


def _text_quality_score(text: str) -> float:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return 0.0

    alpha_count = len(re.findall(r"[A-Za-zА-Яа-яЁё]", cleaned))
    cyr_count = len(re.findall(r"[А-Яа-яЁё]", cleaned))
    long_words = len(re.findall(r"[A-Za-zА-Яа-яЁё]{4,}", cleaned))
    punct_noise = len(re.findall(r"[^A-Za-zА-Яа-яЁё0-9\s.,:;!?()\"'«»\-]", cleaned))

    alpha_ratio = alpha_count / max(1, len(cleaned))
    cyr_ratio = cyr_count / max(1, alpha_count)
    noise_penalty = punct_noise / max(1, len(cleaned))

    return (
        min(len(cleaned), 8000) / 8000.0
        + alpha_ratio * 1.4
        + cyr_ratio * 1.2
        + min(long_words, 200) / 200.0
        - noise_penalty * 3.0
    )


def _normalize_ocr_output(text: str) -> str:
    # Fix common latin/cyrillic look-alike substitutions from OCR.
    translation_table = str.maketrans(
        {
            "A": "А",
            "B": "В",
            "C": "С",
            "E": "Е",
            "H": "Н",
            "K": "К",
            "M": "М",
            "O": "О",
            "P": "Р",
            "T": "Т",
            "X": "Х",
            "Y": "У",
            "a": "а",
            "c": "с",
            "e": "е",
            "o": "о",
            "p": "р",
            "x": "х",
            "y": "у",
        }
    )
    normalized = text.translate(translation_table)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _ocr_pdf_page(pdf_doc, page_index: int) -> str:
    if pdf_doc is None:
        return ""
    try:
        page = pdf_doc.load_page(page_index)
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
        base = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        gray = ImageOps.grayscale(base)
        strong_contrast = ImageOps.autocontrast(gray)
        # Median filter helps suppress scan noise before thresholding.
        denoised = strong_contrast.filter(ImageFilter.MedianFilter(size=3))
        soft_binary = denoised.point(lambda p: 255 if p > 175 else 0)

        variants = [
            (denoised, "--oem 1 --psm 3 -c preserve_interword_spaces=1 -c tessedit_do_invert=0"),
            (denoised, "--oem 1 --psm 4 -c preserve_interword_spaces=1 -c tessedit_do_invert=0"),
            (denoised, "--oem 1 --psm 6 -c preserve_interword_spaces=1 -c tessedit_do_invert=0"),
            (soft_binary, "--oem 1 --psm 6 -c preserve_interword_spaces=1 -c tessedit_do_invert=0"),
            (soft_binary, "--oem 1 --psm 11 -c preserve_interword_spaces=1 -c tessedit_do_invert=0"),
        ]

        best_text = ""
        best_score = -10**9
        for image, config in variants:
            candidate = pytesseract.image_to_string(image, lang=_tesseract_lang(), config=config).strip()
            candidate = _normalize_ocr_output(candidate)
            score = _text_quality_score(candidate)
            if score > best_score:
                best_score = score
                best_text = candidate
        return best_text
    except Exception:
        return ""


def read_document_content(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _read_txt(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return ""


def _parse_toc_file(toc_file: Path) -> dict[int, str]:
    try:
        raw = toc_file.read_text(encoding="utf-8")
    except Exception:
        return {}

    in_toc = False
    entries: list[tuple[int, str]] = []
    for line in raw.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.upper() == "[TOC]":
            in_toc = True
            continue
        if not in_toc:
            continue
        if "|" not in text:
            continue
        left, right = text.split("|", 1)
        try:
            page = int(left.strip())
        except ValueError:
            continue
        title = right.strip()
        if title:
            entries.append((page, title))

    entries.sort(key=lambda item: item[0])
    page_to_chapter: dict[int, str] = {}
    if not entries:
        return page_to_chapter

    # Рубрика действует со страницы из строки ОГЛ до страницы перед следующей строкой (не с p=1).
    # Старый алгоритм присваивал заголовок первой строки страницам 1..(первая_страница_следующей−1),
    # из‑за чего все страницы до «Вместо предисловия» получали этот же chapter.
    # Последняя рубрика должна действовать до конца документа.
    # Мы не знаем финальную страницу на этапе разбора TOC, поэтому задаём большой верхний предел.
    end_of_doc = 100_000
    for i, (start_page, title) in enumerate(entries):
        if i + 1 < len(entries):
            end_inclusive = entries[i + 1][0] - 1
        else:
            end_inclusive = end_of_doc
        for p in range(start_page, end_inclusive + 1):
            page_to_chapter[p] = title
    return page_to_chapter


def _toc_file_candidates_for_doc(doc_name: str, toc_dir: Path) -> list[Path]:
    if not toc_dir.exists():
        return []
    return [
        toc_dir / f"{Path(doc_name).stem}.txt",
        toc_dir / f"{doc_name}.txt",
    ]


def _load_toc_map_for_doc(doc_name: str, toc_dir: Path) -> dict[int, str]:
    for candidate in _toc_file_candidates_for_doc(doc_name, toc_dir):
        if candidate.is_file():
            parsed = _parse_toc_file(candidate)
            if parsed:
                return parsed
    return {}


def _read_toc_file_display_source_doc(toc_path: Path) -> str | None:
    """Читает строку `source_doc:` из `docs_toc/…` до раздела [TOC] (человекочитаемое название книги)."""
    try:
        raw = toc_path.read_text(encoding="utf-8")
    except OSError:
        return None
    for line in raw.splitlines():
        text = line.strip()
        if text.upper() == "[TOC]":
            break
        m = re.match(r"^source_doc:\s*(.+)\s*$", line, re.IGNORECASE)
        if m:
            t = m.group(1).strip()
            if t:
                return t
    return None


def resolve_chapter_for_illustration_page(
    doc_name: str, ill_page: str, default_chapter: str | None, toc_dir: Path
) -> str | None:
    """Глава по номеру страницы и TOC; иначе как в `meta` (для `rag_illustration_chunks`)."""
    p = (ill_page or "").strip()
    if not p.isdigit():
        return (default_chapter or "").strip() or None
    m = _load_toc_map_for_doc(doc_name, toc_dir)
    if m:
        return m.get(int(p), default_chapter) or default_chapter
    return (default_chapter or "").strip() or None


def _apply_toc_map(units: list[TextUnit], toc_map: dict[int, str]) -> list[TextUnit]:
    if not toc_map:
        return units
    mapped: list[TextUnit] = []
    for unit in units:
        chapter = unit.chapter
        if unit.page and unit.page.isdigit():
            chapter = toc_map.get(int(unit.page), chapter)
        mapped.append(TextUnit(text=unit.text, chapter=chapter, page=unit.page))
    return mapped


def build_mapped_text_for_pdf(pdf_path: Path, toc_dir: Path) -> str:
    pages = _read_pdf_pages(pdf_path)
    if not pages:
        return ""

    toc_map = _load_toc_map_for_doc(pdf_path.name, toc_dir)
    blocks: list[str] = [f"source_doc: {pdf_path.name}"]

    first_page = pages[0][0]
    first_chapter = toc_map.get(first_page) if toc_map else None
    if first_chapter:
        blocks.append(f"source_chapter: {first_chapter}")
    blocks.append(f"source_page: {first_page}")
    blocks.append("")

    for page_num, text in pages:
        blocks.append(f"[PAGE {page_num}]")
        chapter = toc_map.get(page_num) if toc_map else None
        if chapter:
            blocks.append(chapter)
        blocks.append(text.strip())
        blocks.append("")

    return "\n".join(blocks).strip()


def get_pdf_page_stats(pdf_path: Path) -> list[tuple[int, int]]:
    pages = _read_pdf_pages(pdf_path)
    return [(page_num, len(text.strip())) for page_num, text in pages]


def read_documents(docs_dir: Path, toc_dir: Path) -> list[IngestDocument]:
    if not docs_dir.exists():
        return []

    documents: list[IngestDocument] = []
    for file_path in docs_dir.rglob("*"):
        if _is_prepare_page_cache_path(file_path):
            continue
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            content = read_document_content(file_path)
        except Exception:
            continue
        if not content.strip():
            continue

        doc_id = str(file_path.relative_to(docs_dir))
        source_doc = _extract_meta(content, "source_doc") or file_path.name
        for tpath in _toc_file_candidates_for_doc(file_path.name, toc_dir):
            if tpath.is_file():
                toc_title = _read_toc_file_display_source_doc(tpath)
                if toc_title:
                    source_doc = toc_title
                break
        source_chapter = _extract_meta(content, "source_chapter")
        source_page = normalize_source_page(_extract_meta(content, "source_page"))
        meta_stripped = _strip_meta_lines(content)
        split_body, split_ill = parse_prepare_split_format(meta_stripped)
        if split_body is not None:
            content_for_index = split_body
            illustration_segments = tuple(split_ill)
        else:
            illustration_segments = tuple(extract_illustration_segments_from_stripped_content(meta_stripped))
            content_for_index = _remove_vl_illustration_description_blocks(meta_stripped)
        content_for_index = _normalize_inline_page_markers(content_for_index)
        if not content_for_index:
            continue
        units = _extract_units(content_for_index, source_chapter, source_page)
        toc_map = _load_toc_map_for_doc(file_path.name, toc_dir)
        units = _apply_toc_map(units, toc_map)
        units = _merge_rubric_only_units(units)
        if not units:
            continue
        documents.append(
            IngestDocument(
                doc_id=doc_id,
                doc_name=file_path.name,
                content=content_for_index,
                meta=SourceMeta(
                    source_doc=source_doc,
                    source_chapter=source_chapter,
                    source_page=source_page,
                ),
                units=units,
                illustration_segments=illustration_segments,
            )
        )
    return documents


def _prefer_chunk_end(text: str, start: int, end: int, chunk_size: int) -> int:
    """Сдвигает конец чанка к последней границе абзаца/строки, чтобы реже резать посередине фразы."""
    if end >= len(text):
        return end
    window = text[start:end]
    min_keep = max(80, chunk_size // 3)
    split_at = window.rfind("\n\n")
    if split_at >= min_keep:
        return start + split_at
    split_at = window.rfind("\n")
    if split_at >= min_keep:
        return start + split_at
    return end


def _split_oversized_unit_text(text: str, chunk_size: int) -> list[str]:
    """Дополнительно дробит длинные юниты по абзацам, чтобы не получать «простыни» в одном чанке."""
    if not text:
        return []
    if chunk_size <= 0 or len(text) <= chunk_size:
        return [text]
    parts: list[str] = []
    paras = re.split(r"\n\s*\n+", text)
    buf = ""
    for p in paras:
        p = p.strip()
        if not p:
            continue
        candidate = (buf + "\n\n" + p).strip() if buf else p
        if len(candidate) <= chunk_size:
            buf = candidate
            continue
        if buf:
            parts.append(buf)
            buf = ""
        if len(p) <= chunk_size:
            buf = p
        else:
            parts.extend(chunk_text(p, chunk_size, 0))
    if buf:
        parts.append(buf)
    return parts if parts else [text]


def _merge_tiny_chunk_records(chunks: list[ChunkRecord], min_len: int) -> list[ChunkRecord]:
    """Склеивает слишком короткие чанки с соседом на той же странице (опционально)."""
    if min_len <= 0 or len(chunks) < 2:
        return chunks
    out: list[ChunkRecord] = []
    acc = chunks[0]
    for nxt in chunks[1:]:
        if (
            len(acc.text) < min_len
            and acc.page == nxt.page
            and (acc.chapter or "") == (nxt.chapter or "")
        ):
            acc = ChunkRecord(
                text=(acc.text + "\n\n" + nxt.text).strip(),
                chapter=acc.chapter,
                page=acc.page,
            )
        else:
            out.append(acc)
            acc = nxt
    out.append(acc)
    return out


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    if not text:
        return []
    if chunk_size <= 0:
        return []
    overlap = max(0, min(overlap, chunk_size - 1))
    chunks: list[str] = []
    start = 0
    while start < len(text):
        raw_end = min(len(text), start + chunk_size)
        end = _prefer_chunk_end(text, start, raw_end, chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        next_start = max(start + 1, end - overlap)
        if next_start >= len(text):
            break
        start = next_start
    return chunks


def chunk_document(
    document: IngestDocument,
    chunk_size: int,
    overlap: int,
    *,
    min_merge_chars: int = 0,
) -> list[ChunkRecord]:
    chunks: list[ChunkRecord] = []
    for unit in document.units:
        for piece in _split_oversized_unit_text(unit.text, chunk_size):
            for chunk in chunk_text(piece, chunk_size, overlap):
                chunks.append(
                    ChunkRecord(
                        text=chunk,
                        chapter=unit.chapter,
                        page=unit.page,
                    )
                )
    if min_merge_chars > 0:
        chunks = _merge_tiny_chunk_records(chunks, min_merge_chars)
    return chunks

