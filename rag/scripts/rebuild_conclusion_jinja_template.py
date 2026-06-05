"""
Пересобирает rag/templates/conclusion_base.docx: Jinja2-плейсхолдеры, подписи полей
в отдельных runs (жирный/курсив), разметка для docxtpl.

Повторный запуск: абзацы, где уже есть «{{» или «{%», не трогаются
(секции источников/исследования — по дополнительным проверкам).
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from rag_service.report_docx_from_template import (  # noqa: E402
    _remove_all_inline_pictures_from_document,
    _insert_photo_after_marker,
)
from rag_service.report_models import ReportPhoto  # noqa: E402

TEMPLATE = ROOT / "templates" / "conclusion_base.docx"
BACKUP = ROOT / "templates" / "conclusion_base_pre_jinja.bak.docx"


def _clear_runs(p) -> None:
    for r in list(p.runs):
        try:
            r._element.getparent().remove(r._element)  # noqa: SLF001
        except Exception:
            pass


def _clear_paragraphs_after_until(
    doc: Document, *, idx: int, stop_prefix: str, max_remove: int = 200
) -> int:
    n = 0
    while n < max_remove:
        body = list(doc.paragraphs)
        if idx + 1 >= len(body):
            break
        nxt = body[idx + 1]
        t2 = (nxt.text or "").strip()
        if t2.startswith(stop_prefix):
            break
        nxt._element.getparent().remove(nxt._element)  # noqa: SLF001
        n += 1
    return n


def _find_index_startswith(doc, prefix: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith(prefix) and "{{" not in (p.text or ""):
            return i
    return None


def _find_index_exact(doc, text: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == text and "for " not in (p.text or ""):
            if "{{" in (p.text or ""):
                continue
            return i
    return None


def _insert_jinja_block_after(  # noqa: PLR0913
    doc: Document, after_paragraph_starts: str, *, line1: str, line2: str, line3: str
) -> bool:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t.startswith(after_paragraph_starts) or "{%" in (p.text or ""):
            continue
        last = p._p  # noqa: SLF001
        for line in (line1, line2, line3):
            new_p = doc.add_paragraph(line)
            last.addnext(new_p._p)  # noqa: SLF001
            last = new_p._p  # noqa: SLF001
        return True
    return False


def _set_para_label_value(
    p,
    label: str,
    key: str,
    *,
    label_bold: bool = True,
    label_italic: bool | None = None,
) -> None:
    if "{{" in (p.text or ""):
        return
    _clear_runs(p)
    r0 = p.add_run(label)
    r0.bold = label_bold
    r0.italic = label_italic
    p.add_run(f" {{{{{key}}}}}")


def _set_para_literal(p, text: str) -> None:
    if "{{" in (p.text or ""):
        return
    _clear_runs(p)
    p.add_run(text)


def _find_para_startswith(doc, startswith: str):
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith(startswith) and "{{" not in (p.text or ""):
            return p
    return None


def main() -> None:
    if not TEMPLATE.is_file():
        raise SystemExit(f"Not found: {TEMPLATE}")

    if not BACKUP.is_file():
        import shutil

        shutil.copy2(TEMPLATE, BACKUP)
        print("Backup saved:", BACKUP)

    doc = Document(str(TEMPLATE))
    nimg = _remove_all_inline_pictures_from_document(doc)
    print("Removed template images:", nimg)

    p = _find_para_startswith(doc, "ЗАКЛЮЧЕНИЕ СПЕЦИАЛИСТА")
    if p is not None:
        _set_para_literal(p, "{{ report_title_line }}")

    p = _find_para_startswith(doc, "Основания производства исследования")
    if p is not None:
        _clear_runs(p)
        r0 = p.add_run("Основания производства исследования: ")
        r0.italic = True
        p.add_run("{{ basis }}")

    p = _find_para_startswith(doc, "Некрасов")
    if p is not None and "{{" not in p.text:
        _set_para_literal(p, "{{ full_name }}")

    p = _find_para_startswith(doc, "Образование:")
    if p is not None:
        _set_para_label_value(p, "Образование:", "education", label_bold=True, label_italic=True)

    p = _find_para_startswith(doc, "Квалификация:")
    if p is not None:
        _set_para_label_value(p, "Квалификация:", "qualification", label_bold=True, label_italic=True)

    p = _find_para_startswith(doc, "Повышение квалификации по")
    if p is not None:
        _set_para_label_value(
            p,
            "Повышение квалификации по дополнительной профессиональной программе: ",
            "additional_training",
            label_bold=True,
            label_italic=True,
        )

    p = _find_para_startswith(doc, "Должность:")
    if p is not None:
        _set_para_label_value(p, "Должность:", "position", label_bold=True, label_italic=True)

    p = _find_para_startswith(doc, "Сфера научных интересов:")
    if p is not None:
        _set_para_label_value(p, "Сфера научных интересов:", "research_interests", label_bold=True, label_italic=True)

    p = _find_para_startswith(doc, "Стаж работы:")
    if p is not None:
        _set_para_label_value(p, "Стаж работы:", "experience_years", label_bold=True, label_italic=True)

    p = _find_para_startswith(doc, "Специалисту")
    if p is not None:
        _set_para_literal(p, "{{ materials_text }}")

    p = _find_para_startswith(doc, "Содержат ли представленные на исследование")
    if p is not None:
        _set_para_literal(p, "{{ question }}")

    # --- источники: удалить пример, вставить jinja-цикл (3 абзаца)
    idx_src = _find_index_startswith(doc, "При производстве исследования использовались")
    if idx_src is not None:
        nxt0 = list(doc.paragraphs)[idx_src + 1] if idx_src + 1 < len(doc.paragraphs) else None
        if nxt0 is not None and (nxt0.text or "").strip().startswith("{%p"):
            pass
        else:
            _clear_paragraphs_after_until(doc, idx=idx_src, stop_prefix="Содержание")
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t.startswith("При производстве исследования использовались") and "{%" not in t:
                    last = p._p  # noqa: SLF001
                    for line in (
                        "{%p for line in source_lines %}",
                        "{{ line|e }}",
                        "{%p endfor %}",
                    ):
                        new_p = doc.add_paragraph(line)
                        last.addnext(new_p._p)  # noqa: SLF001
                        last = new_p._p  # noqa: SLF001
                    break
    elif not _insert_jinja_block_after(
        doc,
        "При производстве исследования использовались",
        line1="{%p for line in source_lines %}",
        line2="{{ line|e }}",
        line3="{%p endfor %}",
    ):
        print("WARNING: jinja for sources: anchor not found or already jinja")

    p = _find_para_startswith(doc, "В качестве базового")
    if p is not None:
        _set_para_literal(p, "{{ methods_text }}")

    idx_iss = _find_index_exact(doc, "Исследование")
    if idx_iss is not None:
        # уже вставляли: после заголовка идёт {%
        nxt = list(doc.paragraphs)[idx_iss + 1] if idx_iss + 1 < len(doc.paragraphs) else None
        if nxt and not (nxt.text or "").strip().startswith("{%p"):
            _clear_paragraphs_after_until(doc, idx=idx_iss, stop_prefix="Вывод")
            anchor_p = list(doc.paragraphs)[idx_iss]
            last = anchor_p._p  # noqa: SLF001
            for line in (
                "{%p for rtext in research_paragraphs %}",
                "{{ rtext|e }}",
                "{%p endfor %}",
            ):
                new_p = doc.add_paragraph(line)
                last.addnext(new_p._p)  # noqa: SLF001
                last = new_p._p  # noqa: SLF001
            m = doc.add_paragraph("(Фото № 1)")
            last.addnext(m._p)  # noqa: SLF001

    p = _find_para_startswith(doc, "Представленный на")
    if p is not None and "conclusion_text" not in (p.text or ""):
        _set_para_literal(p, "{{ conclusion_text }}")

    p = _find_para_startswith(doc, "Примечание:")
    if p is not None and "note_text" not in (p.text or ""):
        _clear_runs(p)
        r0 = p.add_run("Примечание: ")
        r0.italic = True
        p.add_run("{{ note_text }}")

    # Сохраняем **нераскрытый** шаблон (только Jinja, без render)
    doc.save(str(TEMPLATE))
    print("Written template with Jinja2 placeholders:", TEMPLATE)

    # Дымовой: render в отдельный файл, не портя шаблон
    from docxtpl import DocxTemplate  # type: ignore[import-not-found]  # noqa: WPS433

    tpl2 = DocxTemplate(str(TEMPLATE))
    ctx = {
        "report_title_line": "T",
        "basis": "b",
        "full_name": "f",
        "education": "e",
        "qualification": "q",
        "additional_training": "a",
        "position": "p",
        "research_interests": "r",
        "experience_years": "x",
        "materials_text": "m",
        "question": "Q",
        "source_lines": ["1. s"],
        "methods_text": "M",
        "research_paragraphs": ["R1", "R2"],
        "conclusion_text": "C",
        "note_text": "N",
    }
    tpl2.render(ctx)
    b3 = io.BytesIO()
    tpl2.save(b3)
    b3.seek(0)
    d3 = Document(b3)
    _remove_all_inline_pictures_from_document(d3)
    imgp = Path(__file__).resolve().parents[2] / "exemple" / "photo1.jpg"
    if imgp.is_file():
        _insert_photo_after_marker(d3, "(Фото № 1)", ReportPhoto(1, imgp.read_bytes(), "image/jpeg"))
    else:
        print("Skip photo in smoke: exemple/photo1.jpg not found")
    d3.save(ROOT / "templates" / "conclusion_jinja_smoke_test.docx")
    print("Smoke render written: templates/conclusion_jinja_smoke_test.docx")


if __name__ == "__main__":
    main()
